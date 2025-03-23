import os
import json
import re
import random
import string
from docx import Document
import shutil

def generate_random_name():
    """랜덤 이름 생성 (한글)"""
    first_names = ['김', '이', '박', '최', '정', '강', '조', '윤', '장', '임', '한', '오', '서', '신', '권', '황', '안', '송', '전', '홍']
    last_names = ['준', '민', '서', '지', '예', '원', '현', '수', '영', '재', '주', '성', '우', '진', '선', '은', '혜', '정', '윤', '태', '동', '민', '석', '준', '윤', '호', '현', '아', '종', '훈']
    
    first = random.choice(first_names)
    last = random.choice(last_names) + random.choice(last_names)
    
    return first + last

def generate_random_student_id():
    """랜덤 학번 생성 (ex: 202312345)"""
    year = random.randint(2018, 2023)
    number = random.randint(10000, 99999)
    return f"{year}{number}"

def anonymize_docx_files():
    """docx 파일 내의 이름과 학번을 익명화하고 파일명도 변경"""
    
    # 'out' 폴더 경로 확인
    out_folder = 'out'
    if not os.path.exists(out_folder):
        print(f"'{out_folder}' 폴더가 존재하지 않습니다.")
        return
    
    # 결과물을 저장할 'anonymized' 폴더 생성
    anonymized_folder = 'anonymized'
    if not os.path.exists(anonymized_folder):
        os.makedirs(anonymized_folder)
    
    # 맵핑 정보를 저장할 딕셔너리
    mapping = {}
    
    # 'out' 폴더에 있는 모든 docx 파일 처리
    for filename in os.listdir(out_folder):
        if filename.endswith('.docx'):
            file_path = os.path.join(out_folder, filename)
            
            # 파일명에서 이름과 학번 추출 (파일명 형식은 "이름_학번.docx"이라고 가정)
            name_student_id = os.path.splitext(filename)[0]
            parts = name_student_id.split('_')
            
            if len(parts) >= 2:
                original_name = parts[0]
                original_student_id = parts[1]
                
                # 새로운 랜덤 이름과 학번 생성
                new_name = generate_random_name()
                new_student_id = generate_random_student_id()
                
                # 맵핑 정보 저장
                mapping[original_name + "_" + original_student_id] = {
                    "original_name": original_name,
                    "original_student_id": original_student_id,
                    "new_name": new_name,
                    "new_student_id": new_student_id
                }
                
                # docx 파일 내용 수정
                doc = Document(file_path)
                
                # 모든 단락에서 이름과 학번 대체
                for paragraph in doc.paragraphs:
                    if original_name in paragraph.text:
                        paragraph.text = paragraph.text.replace(original_name, new_name)
                    if original_student_id in paragraph.text:
                        paragraph.text = paragraph.text.replace(original_student_id, new_student_id)
                
                # 테이블 내용 대체
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if original_name in cell.text:
                                cell.text = cell.text.replace(original_name, new_name)
                            if original_student_id in cell.text:
                                cell.text = cell.text.replace(original_student_id, new_student_id)
                
                # 새로운 파일명으로 저장
                new_filename = f"{new_name}_{new_student_id}.docx"
                new_file_path = os.path.join(anonymized_folder, new_filename)
                doc.save(new_file_path)
                
                print(f"익명화 완료: {filename} -> {new_filename}")
            else:
                # 파일명 형식이 예상과 다를 경우 원본 파일 복사
                print(f"경고: 파일명 형식이 예상과 다릅니다: {filename}")
                shutil.copy2(file_path, os.path.join(anonymized_folder, filename))
    
    # 맵핑 정보를 JSON 파일로 저장
    with open('name_student_id_mapping.json', 'w', encoding='utf-8') as f:
        json.dump(mapping, f, ensure_ascii=False, indent=4)
    
    print(f"맵핑 정보가 'name_student_id_mapping.json' 파일에 저장되었습니다.")

def restore_docx_files():
    """익명화된 docx 파일을 원래 정보로 복구"""
    
    # 'anonymized' 폴더 경로 확인
    anonymized_folder = 'anonymized'
    if not os.path.exists(anonymized_folder):
        print(f"'{anonymized_folder}' 폴더가 존재하지 않습니다.")
        return
    
    # 결과물을 저장할 'restored' 폴더 생성
    restored_folder = 'restored'
    if not os.path.exists(restored_folder):
        os.makedirs(restored_folder)
    
    # 맵핑 정보 로드
    try:
        with open('name_student_id_mapping.json', 'r', encoding='utf-8') as f:
            mapping = json.load(f)
    except FileNotFoundError:
        print("맵핑 정보 파일을 찾을 수 없습니다.")
        return
    
    # 역방향 맵핑 생성 (새 이름+학번 -> 원래 정보)
    reverse_mapping = {}
    for key, value in mapping.items():
        new_key = value["new_name"] + "_" + value["new_student_id"]
        reverse_mapping[new_key] = {
            "original_name": value["original_name"],
            "original_student_id": value["original_student_id"]
        }
    
    # 'anonymized' 폴더에 있는 모든 docx 파일 처리
    for filename in os.listdir(anonymized_folder):
        if filename.endswith('.docx'):
            file_path = os.path.join(anonymized_folder, filename)
            
            # 파일명에서 이름과 학번 추출
            name_student_id = os.path.splitext(filename)[0]
            parts = name_student_id.split('_')
            
            if len(parts) >= 2:
                new_name = parts[0]
                new_student_id = parts[1]
                new_key = new_name + "_" + new_student_id
                
                if new_key in reverse_mapping:
                    # 원래 정보 가져오기
                    original_info = reverse_mapping[new_key]
                    original_name = original_info["original_name"]
                    original_student_id = original_info["original_student_id"]
                    
                    # docx 파일 내용 수정
                    doc = Document(file_path)
                    
                    # 모든 단락에서 익명화된 이름과 학번을 원래 정보로 대체
                    for paragraph in doc.paragraphs:
                        if new_name in paragraph.text:
                            paragraph.text = paragraph.text.replace(new_name, original_name)
                        if new_student_id in paragraph.text:
                            paragraph.text = paragraph.text.replace(new_student_id, original_student_id)
                    
                    # 테이블 내용 대체
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if new_name in cell.text:
                                    cell.text = cell.text.replace(new_name, original_name)
                                if new_student_id in cell.text:
                                    cell.text = cell.text.replace(new_student_id, original_student_id)
                    
                    # 원래 파일명으로 저장
                    original_filename = f"{original_name}_{original_student_id}.docx"
                    restored_file_path = os.path.join(restored_folder, original_filename)
                    doc.save(restored_file_path)
                    
                    print(f"복구 완료: {filename} -> {original_filename}")
                else:
                    # 맵핑 정보에 없는 경우 원본 파일 복사
                    print(f"경고: 맵핑 정보에 없는 파일: {filename}")
                    shutil.copy2(file_path, os.path.join(restored_folder, filename))
            else:
                # 파일명 형식이 예상과 다를 경우 원본 파일 복사
                print(f"경고: 파일명 형식이 예상과 다릅니다: {filename}")
                shutil.copy2(file_path, os.path.join(restored_folder, filename))
    
    print(f"모든 파일이 '{restored_folder}' 폴더에 복구되었습니다.")

# 메인 코드 실행
if __name__ == "__main__":
    print("1. 개인정보 익명화")
    print("2. 원래 정보로 복구")
    choice = input("작업을 선택하세요 (1 또는 2): ")
    
    if choice == "1":
        anonymize_docx_files()
    elif choice == "2":
        restore_docx_files()
    else:
        print("올바른 선택이 아닙니다. 1 또는 2를 입력하세요.")